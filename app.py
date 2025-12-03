# app.py

import tempfile
from pathlib import Path

import streamlit as st
from pdf2docx import Converter
from docx import Document
from deep_translator import GoogleTranslator
from docx2pdf import convert as docx2pdf_convert


def pdf_to_docx(
    pdf_path: str,
    docx_path: str,
    start_page: int | None = None,
    end_page: int | None = None,
    multi_processing: bool = False,
):
    cv = Converter(pdf_path)

    # pdf2docx uses zero-based page indices
    convert_kwargs = {}
    if start_page is not None:
        convert_kwargs["start"] = max(start_page - 1, 0)
    if end_page is not None:
        # end is inclusive index in pdf2docx
        convert_kwargs["end"] = max(end_page - 1, 0)

    cv.convert(docx_path, multi_processing=multi_processing, **convert_kwargs)
    cv.close()


def translate_docx_by_paragraph(
    input_path: str,
    output_path: str,
    source_lang: str = "auto",
    target_lang: str = "zh-TW",
):
    """
    Translate a .docx file paragraph-by-paragraph.
    Paragraph-level formatting is preserved (e.g. headings),
    but inline formatting (bold/italic on specific words) may be lost.
    """
    input_path = Path(input_path)
    output_path = Path(output_path)

    doc = Document(str(input_path))
    translator = GoogleTranslator(source=source_lang, target=target_lang)

    cache: dict[str, str] = {}

    def safe_translate(text: str) -> str:
        stripped = text.strip()
        if not stripped:
            return text  # keep empty/whitespace-only paragraphs as-is

        if stripped in cache:
            translated = cache[stripped]
        else:
            try:
                translated = translator.translate(stripped)
            except Exception as e:
                # Fallback: keep original text if translation fails
                print(f"Translation error for text: {stripped[:50]!r}... -> {e}")
                return text

            if not translated:
                # If API returns None or empty, keep original
                return text

            cache[stripped] = translated

        # Preserve leading/trailing spaces
        leading = len(text) - len(text.lstrip(" "))
        trailing = len(text) - len(text.rstrip(" "))
        return " " * leading + translated + " " * trailing

    # Translate normal paragraphs
    for para in doc.paragraphs:
        para.text = safe_translate(para.text)

    # Translate paragraphs inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.text = safe_translate(para.text)

    doc.save(str(output_path))


def main():
    st.title("PDF → DOCX Converter (with optional translation)")

    st.markdown(
        "Upload a PDF, optionally translate the converted Word document, "
        "and optionally convert the translated document to PDF."
    )

    uploaded_file = st.file_uploader(
        "Upload PDF file",
        type=["pdf"],
        accept_multiple_files=False,
    )

    # ---- Conversion options ----
    st.subheader("Conversion options")

    col1, col2 = st.columns(2)
    with col1:
        start_page = st.number_input(
            "Start page (optional)",
            min_value=1,
            value=1,
            step=1,
        )
        use_start = st.checkbox("Use start page", value=False)
    with col2:
        end_page = st.number_input(
            "End page (Not inclusive, optional)",
            min_value=1,
            value=1,
            step=1,
        )
        use_end = st.checkbox("Use end page", value=False)

    multi_processing = st.checkbox(
        "Enable multi-processing (faster on large PDFs)", value=False
    )

    # ---- Translation options ----
    st.subheader("Translation options")

    translate_enabled = st.checkbox("Translate output DOCX", value=False)

    # Load supported languages from deep-translator
    langs_dict = GoogleTranslator().get_supported_languages(as_dict=True)
    # langs_dict: {language_name: language_code}
    language_names = sorted(langs_dict.keys())

    # Try to set default target to Traditional Chinese if available, else first language.
    default_lang_name = None
    for name, code in langs_dict.items():
        if code.lower() in ("zh-tw", "zh-tw".lower()):
            default_lang_name = name
            break
    if default_lang_name is None and language_names:
        default_lang_name = language_names[0]

    target_lang_name = None
    target_lang_code = None

    if translate_enabled:
        target_lang_name = st.selectbox(
            "Target language",
            language_names,
            index=language_names.index(default_lang_name) if default_lang_name in language_names else 0,
        )
        target_lang_code = langs_dict[target_lang_name]

    convert_translated_to_pdf = st.checkbox(
        "Convert translated DOCX to PDF",
        value=False,
        help="Only applies if translation is enabled.",
    )

    # ---- Output filename ----
    default_output_name = "output"
    
    if uploaded_file is not None:
        name_stem = Path(uploaded_file.name).stem
        if convert_translated_to_pdf:
            default_output_name = f"{name_stem}.pdf"
        else:
            default_output_name = f"{name_stem}.docx"

    output_filename = st.text_input(
        "Base output filename", value=default_output_name
    )

    # ---- Convert button ----
    convert_btn = st.button("Convert", disabled=(uploaded_file is None))

    if convert_btn and uploaded_file is not None:
        # Ensure .docx suffix for DOCX output name
        if not output_filename.lower().endswith(".docx"):
            output_filename += ".docx"

        # Ensure page range is consistent
        sp = start_page if use_start else None
        ep = end_page if use_end else None
        if sp is not None and ep is not None and ep < sp:
            st.error("End page cannot be smaller than start page.")
            return

        with st.spinner("Processing..."):
            # 1. Write uploaded PDF to a temporary file
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_pdf:
                tmp_pdf.write(uploaded_file.read())
                tmp_pdf_path = tmp_pdf.name

            # 2. Prepare temp output DOCX path for initial conversion
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_docx:
                tmp_docx_path = tmp_docx.name

            # 3. PDF -> DOCX
            pdf_to_docx(
                pdf_path=tmp_pdf_path,
                docx_path=tmp_docx_path,
                start_page=sp,
                end_page=ep,
                multi_processing=multi_processing,
            )

            final_docx_path = tmp_docx_path

            # 4. Optional translation (DOCX -> translated DOCX)
            if translate_enabled and target_lang_code is not None:
                with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_translated_docx:
                    tmp_translated_docx_path = tmp_translated_docx.name

                translate_docx_by_paragraph(
                    input_path=final_docx_path,
                    output_path=tmp_translated_docx_path,
                    source_lang="auto",
                    target_lang=target_lang_code,
                )
                final_docx_path = tmp_translated_docx_path

            # 5. Optional conversion of translated DOCX to PDF
            if translate_enabled and convert_translated_to_pdf:
                with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_pdf_out:
                    tmp_pdf_out_path = tmp_pdf_out.name

                # DOCX -> PDF
                docx2pdf_convert(final_docx_path, tmp_pdf_out_path)

                with open(tmp_pdf_out_path, "rb") as f:
                    file_bytes = f.read()

                download_filename = Path(output_filename).with_suffix(".pdf").name
                mime = "application/pdf"
                label = "Download translated PDF"
            else:
                # Default: just serve DOCX (translated if enabled, otherwise plain)
                with open(final_docx_path, "rb") as f:
                    file_bytes = f.read()

                # If translated but not PDF, you may want to tag name, e.g. "_translated"
                if translate_enabled:
                    base = Path(output_filename).stem
                    download_filename = f"{base}_translated.docx"
                else:
                    download_filename = output_filename

                mime = (
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                label = "Download DOCX" if not translate_enabled else "Download translated DOCX"

        st.success("Conversion complete.")
        st.download_button(
            label=label,
            data=file_bytes,
            file_name=download_filename,
            mime=mime,
        )


if __name__ == "__main__":
    main()
